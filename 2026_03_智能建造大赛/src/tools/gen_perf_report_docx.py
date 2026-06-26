#!/usr/bin/env python3
"""
gen_perf_report_docx.py · 把"性能报告_模板.md"渲染成结构化 Word 性能评估报告初稿(交付物③)

为什么:官方交付物③要求一份性能评估报告(完整 PR 曲线 + 各类目标检测精度 +
  指定边缘端真实推理 FPS≥30),这是 30FPS 硬门槛**唯一可信的证明**。评委靠它打
  "完成度 + 工程严谨性"。本脚本仿照 gen_tech_plan_docx.py 的 python-docx 渲染管线,
  把模板落成带:
    - 封面 + 硬指标结论速览表(开篇焊死出局线)
    - 六个章节(对应模板 §1-§6):实验环境/检测精度/边缘端实时性/消融/系统闭环/结论
    - 内嵌已生成的 output/figs/ 三张性能图:分桶召回 / 各类 PR 曲线 / FPS-精度权衡
  生成 .docx 到 output/docx/。**这是"初稿"**:模板中所有真值留 ___ 待 eval.py /
  trt_infer_orin.py 产出后回填(PLACEHOLDER 真值),表格结构与模板逐表同构,方便直填。

  设计原则:表头/行结构与 性能报告_模板.md 严格对齐(同构表格),空真值统一用 ___ 占位,
  方便后续把 eval/trt 真值直接填进同构表格。全文匿名(无校名/logo)。

依赖: python-docx。无 GPU。
用法:
  python3 gen_perf_report_docx.py            # 生成 docx
  python3 gen_perf_report_docx.py --selftest # 自测(生成到临时目录并校验结构/非空/内嵌图)
"""
import argparse
import os
import sys
from pathlib import Path

SRC_DIR = Path(__file__).resolve().parent.parent
PROJ_DIR = SRC_DIR.parent
OUT_DIR = PROJ_DIR / "output" / "docx"
FIG_DIR = PROJ_DIR / "output" / "figs"

PH = "___"  # PLACEHOLDER 真值占位符(待 eval.py / trt_infer_orin.py 回填)

# 本报告会尝试内嵌的三张性能图(与官方交付物③要求一一对应)
PERF_FIGS = [
    "report_fig1_bucket_recall.png",  # 分桶召回(§2 小目标不漏)
    "report_fig2_pr_curves.png",      # 各类 PR 曲线(§2 核心交付)
    "report_fig3_fps_accuracy.png",   # FPS-精度权衡 + 30FPS 红线(§3)
]


def _set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold


def _add_table(doc, header, rows, col_widths=None):
    from docx.shared import Inches
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        _set_cell_text(t.rows[0].cells[i], h, bold=True)
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            _set_cell_text(cells[i], v)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t


def _maybe_add_figure(doc, fig_name, caption):
    """若图存在则插入(居中,宽 6 寸)并加题注;不存在则插入文字占位。返回是否内嵌。"""
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    fig_path = FIG_DIR / fig_name
    if fig_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(fig_path), width=Inches(6.0))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        return True
    else:
        doc.add_paragraph(f"[配图占位: {fig_name} — 运行 gen_report_figs.py 生成] {caption}")
        return False


def build_doc():
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 中文默认字体
    style = doc.styles["Normal"]
    style.font.name = "Songti SC"
    style.font.size = Pt(10.5)
    try:
        from docx.oxml.ns import qn
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    except Exception:
        pass

    # ---- 封面 ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("低空智能水上救援装备\n高精度视觉识别 · 性能评估报告")
    r.bold = True; r.font.size = Pt(22)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("中国研究生智能建造大赛(2026)· 揭榜赛道 赛题 #7\n"
                     "基于边缘计算的低空智能水上救援装备高精度视觉识别技术")
    rs.font.size = Pt(11); rs.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"性能评估报告(初稿) · 匿名版 · 真值待回填({PH})\n"
                 "[交付物③:完整 PR 曲线 + 各类检测精度 + 边缘端真实推理 FPS ≥30]").font.size = Pt(9)
    doc.add_page_break()

    # ---- 0. 结论速览 · 硬指标对照 ----
    doc.add_heading("0. 结论速览 · 硬指标对照", level=1)
    doc.add_paragraph(
        "本报告以官方交付物③的三类硬证据为骨架:各类 Precision-Recall 曲线、各类目标检测精度(mAP),"
        f"以及**指定边缘端真实推理 FPS**(出局线 ≥30FPS)。报告原则:30FPS 必须报"
        "「含画框+编码」的端到端值,而非裸推理;数据由 eval.py / trt_infer_orin.py 产出后填入。"
        f"所有真值在初稿阶段统一以 {PH} 占位,结构与本队同构表格对齐,可直接回填。")
    _add_table(doc,
               ["硬指标", "官方门槛", "本方案达成", "证据出处"],
               [["边缘端实时性", "动态视频 ≥30FPS", f"{PH} FPS(含编码端到端)", "本报告 §3 三档 FPS"],
                ["检测精度(mAP)", "各类目标检测精度", f"mAP={PH}", "本报告 §2 精度表"],
                ["完整 PR 曲线", "各类 PR 曲线", "落水人员/船只/浮标", "本报告 §2 PR 曲线图"],
                ["小目标召回", "小目标不漏检", f"small召回={PH}", "本报告 §2 分桶召回图"]],
               col_widths=[1.4, 1.7, 1.8, 1.6])
    doc.add_paragraph(
        f"一句话结论(待回填):端到端含编码 {PH} FPS ≥ 30 ✅;mAP {PH};小目标召回 {PH};"
        f"最优工作点 imgsz {PH} + {PH} 精度模式。")

    # ---- 1. 实验环境(对应模板 §1) ----
    doc.add_heading("1. 实验环境", level=1)
    doc.add_paragraph("训练端与机载边缘端配置分列;端侧务必声明 Super Mode 与供电方式(直接影响 FPS 复现)。")
    _add_table(doc,
               ["项", "配置"],
               [["训练", f"RTX4070S, ultralytics {PH}, torch {PH}, imgsz 1024"],
                ["端侧", f"Jetson Orin Nano 8G, JetPack 6.2, TensorRT {PH}, "
                        "Super Mode(MAXN_SUPER 25W)+jetson_clocks, DC 供电"],
                ["数据", f"SeaDronesSee+AFO(公开)/ 命题方二开数据({PH}张, 训/验/测 {PH})"]],
               col_widths=[1.2, 5.6])

    # ---- 2. 检测精度(对应模板 §2) ----
    doc.add_heading("2. 检测精度(各类 PR 曲线 + mAP)", level=1)
    doc.add_paragraph("各类目标检测精度;救援场景以落水人员的 recall 为第一优先级。")
    _add_table(doc,
               ["类别", "AP@0.5", "AP@0.5:0.95", "备注"],
               [["落水人员", PH, PH, "救援核心,recall 优先"],
                ["船只", PH, PH, ""],
                ["浮标", PH, PH, ""],
                ["mAP", PH, PH, "三类平均"]],
               col_widths=[1.6, 1.3, 1.6, 2.3])
    doc.add_paragraph("交付图(已生成占位曲线,真值回填后替换):", style="List Bullet")
    _maybe_add_figure(doc, "report_fig2_pr_curves.png",
                      "图2-1 各类 Precision-Recall 曲线(占位值)。救援场景 recall 优先,"
                      "落水人员高 recall 区为验收重点。")
    doc.add_paragraph(
        "按目标像素面积分桶(tiny/small/medium/large)的召回——证明小目标不漏检"
        "(eval.py 分桶召回输出):", style="List Bullet")
    _maybe_add_figure(doc, "report_fig1_bucket_recall.png",
                      "图2-2 按目标像素面积分桶的召回(占位值)。P2+NWD+GT-Glint 链路在 "
                      "tiny/small 带来最大增益,证明小目标不漏检。")

    # ---- 3. 边缘端实时性(对应模板 §3,30FPS 硬门槛证明) ----
    doc.add_heading("3. 边缘端实时性(30FPS 硬门槛证明)", level=1)
    doc.add_paragraph(
        "报端到端三档(trt_infer_orin.py 输出),以「③含编码」为准。"
        "Orin Nano 无硬件 NVENC,软编码 H265 已计入端到端时延。")
    _add_table(doc,
               ["档位", "imgsz", "精度模式", "时延ms", "FPS", "vs 30"],
               [["① 裸推理", "640", "INT8", PH, PH, PH],
                ["② 含后处理(解码+NMS+画框)", "640", "INT8", PH, PH, PH],
                ["③ 含编码(实战)", "640", "INT8", PH, PH, "✅/❌"]],
               col_widths=[2.4, 0.8, 1.0, 0.9, 0.7, 0.8])
    for s in ["FP16 vs INT8 × 640 vs 768 的 FPS-召回帕累托曲线(选最优工作点);",
              "tegrastats 记录:功耗(W)、显存(MB)、nvpmodel 档位。"]:
        doc.add_paragraph(s, style="List Bullet")
    _maybe_add_figure(doc, "report_fig3_fps_accuracy.png",
                      "图3-1 FPS-精度权衡 + 30FPS 红线 + 帕累托前沿(占位值),用于选定端侧最优工作点。")

    # ---- 4. 消融实验(对应模板 §4) ----
    doc.add_heading("4. 消融实验(技术创新性的硬证据,逐项独立 A/B)", level=1)
    doc.add_paragraph("每行只改一项,其余固定;禁止一次叠多项凭感觉。")
    _add_table(doc,
               ["配置", "mAP", "small召回", "端侧FPS", "结论"],
               [["baseline (yolo12n)", PH, PH, PH, "基线"],
                ["+ P2 头", PH, PH, PH, "小目标↑"],
                ["+ NWD 标签分配", PH, PH, PH, "位置鲁棒↑"],
                ["+ Wise-IoU", PH, PH, PH, "回归↑"],
                ["+ 水面物理增广 + GT-Glint 难负样本", PH, PH, PH, "反光误报↓"],
                ["去 P5(四头→三头)", PH, PH, PH, "帧率↑/精度?"],
                ["+ INT8 量化", PH, PH, PH, "速度↑/掉点?"]],
               col_widths=[2.8, 0.9, 1.1, 1.0, 1.0])

    # ---- 5. 系统闭环验证(对应模板 §5) ----
    doc.add_heading("5. 系统闭环验证", level=1)
    for s in ["检测框→GPS:用已知 GPS 浮标做真值校验,报投影误差区间(诚实,几十米级);",
              "QGroundControl 融合 demo:录屏 + 现场实时双轨;",
              "时序滤波前后误报对比(反光闪点抑制率)。"]:
        doc.add_paragraph(s, style="List Bullet")
    _add_table(doc,
               ["闭环项", "验证方式", "指标", "结果"],
               [["检测框→GPS", "已知 GPS 浮标真值", "投影误差区间(m)", PH],
                ["QGC 视频流融合", "录屏 + 现场双轨", "端到端可视化时延(ms)", PH],
                ["时序滤波误报抑制", "连续≥k帧确认 前/后", "反光闪点误报率(%)", PH]],
               col_widths=[1.8, 1.9, 2.0, 1.1])

    # ---- 6. 结论(对应模板 §6) ----
    doc.add_heading("6. 结论", level=1)
    for s in [f"硬指标:端到端含编码 {PH} FPS ≥ 30 ✅;mAP {PH};小目标召回 {PH};",
              f"最优工作点:imgsz {PH} + {PH} 精度模式;",
              "完成度证据:本报告 §2 各类 PR 曲线 + §3 三档端到端 FPS 共同构成 30FPS 出局线的可信证明。"]:
        doc.add_paragraph(s, style="List Bullet")
    doc.add_paragraph(
        f"回填须知:本报告为初稿,所有 {PH} 为待回填真值,来源为 eval.py(精度/分桶召回)"
        "与 trt_infer_orin.py(端侧三档 FPS / tegrastats)。回填时仅替换 ___,表格结构保持不变,"
        "并用真值曲线图替换 output/figs 下的占位图。")

    return doc


def generate(out_dir=OUT_DIR):
    out_dir = Path(out_dir); out_dir.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    out_path = out_dir / "性能报告_初稿.docx"
    doc.save(str(out_path))
    return out_path


def _selftest():
    import tempfile
    import zipfile
    from docx import Document
    ok = True

    def check(c, m):
        nonlocal ok
        print(("  OK  " if c else "  XX  ") + m)
        ok = ok and c

    figs_present = [f for f in PERF_FIGS if (FIG_DIR / f).exists()]

    with tempfile.TemporaryDirectory() as td:
        out = generate(Path(td))
        check(out.exists() and out.stat().st_size > 8000, f"docx 生成且非空({out.stat().st_size}B)")

        # 合法 OOXML: .docx 应为合法 zip 且含 word/document.xml
        check(zipfile.is_zipfile(str(out)), "docx 为合法 zip(OOXML 容器)")
        with zipfile.ZipFile(str(out)) as z:
            names = z.namelist()
            bad = z.testzip()
            check(bad is None, "zip 内部无损坏(testzip 通过)")
            check("word/document.xml" in names, "含 word/document.xml(合法 OOXML 主文档)")
            check("[Content_Types].xml" in names, "含 [Content_Types].xml(OOXML 必备)")
            media = [n for n in names if n.startswith("word/media/")]
            check(len(media) == len(figs_present),
                  f"内嵌图片数 {len(media)} == 现存性能图数 {len(figs_present)}")

        # 重新打开校验段落/表格/章节
        d = Document(str(out))
        paras = [p.text for p in d.paragraphs]
        check(len(paras) >= 30, f"段落数 {len(paras)} (>=30)")
        headings = [p.text for p in d.paragraphs if p.style.name.startswith("Heading")]
        check(len(headings) >= 6, f"章节标题数 {len(headings)} (含 0-6 章)")
        for key in ["实验环境", "检测精度", "实时性", "消融", "闭环", "结论"]:
            check(any(key in h for h in headings), f"含章节关键词「{key}」")

        # 表格:速览/环境/精度/FPS三档/消融/闭环 共 6 张
        check(len(d.tables) >= 6, f"含表格 {len(d.tables)} 张(速览/环境/精度/FPS/消融/闭环)")

        # 精度表(§2)表头与模板同构
        prec = [t for t in d.tables if [c.text for c in t.rows[0].cells][:3] == ["类别", "AP@0.5", "AP@0.5:0.95"]]
        check(len(prec) >= 1, "精度表表头与模板同构(类别/AP@0.5/AP@0.5:0.95/备注)")
        if prec:
            check(len(prec[0].rows) >= 5, f"精度表含 4 类行({len(prec[0].rows)-1} 数据行)")

        # FPS 三档表(§3)
        fps = [t for t in d.tables if [c.text for c in t.rows[0].cells][:2] == ["档位", "imgsz"]]
        check(len(fps) >= 1 and len(fps[0].rows) >= 4, "FPS 三档表含 3 档行 + 表头")
        if fps:
            check("vs 30" in [c.text for c in fps[0].rows[0].cells], "FPS 表含「vs 30」出局线列")

        # 消融表 7 行配置
        ablation = [t for t in d.tables if t.rows[0].cells[0].text == "配置"]
        check(len(ablation) >= 1 and len(ablation[0].rows) >= 8, "消融表含 >=7 行配置")

        # PLACEHOLDER 真值: 全文应仍含 ___ 占位
        full_text = "\n".join(paras) + "\n".join(
            c.text for t in d.tables for row in t.rows for c in row.cells)
        check(PH in full_text, "全文保留 ___ PLACEHOLDER 真值待回填")

    print(f"\n  (内嵌性能图: {len(figs_present)}/{len(PERF_FIGS)} 张就绪)")
    print("\n" + ("OK 性能报告 docx 生成器 自测通过" if ok else "XX 自测未通过"))
    return ok


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--selftest", action="store_true")
    ap.add_argument("--out", default=str(OUT_DIR))
    a = ap.parse_args()
    if a.selftest:
        sys.exit(0 if _selftest() else 1)
    out = generate(a.out)
    figs_in = sum(1 for f in PERF_FIGS if (FIG_DIR / f).exists())
    print(f"[OK] 生成 {out}  ({os.path.getsize(out)} bytes, 内嵌 {figs_in}/{len(PERF_FIGS)} 张性能图)")


if __name__ == "__main__":
    main()
