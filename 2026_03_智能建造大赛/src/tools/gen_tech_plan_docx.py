#!/usr/bin/env python3
"""
gen_tech_plan_docx.py · 把"技术方案_大纲.md"渲染成结构化 Word 技术方案初稿(交付物①)

为什么:官方交付物①要求一份技术方案文档(选型依据/小目标特征/注意力/轻量化四要点),
  评委靠它打"技术创新性 + 完成度"。本脚本用 python-docx 把大纲落成带:
    - 封面 + 硬指标达标对照表(开篇焊死出局线)
    - 八个章节(对应大纲 §1-§8),含模型选型对照表、消融实验表
    - 若已生成报告图表(output/figs/report_fig*.png),自动插入对应章节作占位配图
  生成 .docx 到 output/docx/。**这是"初稿"**:留 ___ 待真数据/真权重回填。

  设计原则:章节文字来自大纲的技术判断(非空话),表格结构与 性能报告_模板.md 对齐,
  方便后续把 eval/trt 真值直接填进同构表格。全文匿名(无校名/logo)。

依赖: python-docx。无 GPU。
用法:
  python3 gen_tech_plan_docx.py            # 生成 docx
  python3 gen_tech_plan_docx.py --selftest # 自测(生成到临时目录并校验结构/非空)
"""
import argparse
import os
import sys
from pathlib import Path

SRC_DIR = Path(__file__).resolve().parent.parent
PROJ_DIR = SRC_DIR.parent
OUT_DIR = PROJ_DIR / "output" / "docx"
FIG_DIR = PROJ_DIR / "output" / "figs"


def _set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold


def _add_table(doc, header, rows, col_widths=None):
    from docx.shared import Inches
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        _set_cell_text(t.rows[0].cells[i], h, bold=True)
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            _set_cell_text(cells[i], v)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t


def _maybe_add_figure(doc, fig_name, caption):
    """若图存在则插入(居中,宽 6 寸)并加题注;不存在则插入文字占位。"""
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    fig_path = FIG_DIR / fig_name
    if fig_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(fig_path), width=Inches(6.0))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        from docx.shared import Pt
        r.font.size = Pt(9)
        return True
    else:
        doc.add_paragraph(f"[配图占位: {fig_name} — 运行 gen_report_figs.py 生成] {caption}")
        return False


def build_doc():
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 中文默认字体
    style = doc.styles["Normal"]
    style.font.name = "Songti SC"
    style.font.size = Pt(10.5)
    try:
        from docx.oxml.ns import qn
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    except Exception:
        pass

    # ---- 封面 ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("低空智能水上救援装备\n高精度视觉识别技术方案")
    r.bold = True; r.font.size = Pt(22)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("中国研究生智能建造大赛(2026)· 揭榜赛道 赛题 #7\n"
                     "基于边缘计算的低空智能水上救援装备高精度视觉识别技术")
    rs.font.size = Pt(11); rs.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("技术方案文档(初稿) · 匿名版 · ___ 张数据 / 真值待回填\n"
                 "[交付物①:模型选型 + 海上小目标特征 + 注意力增强 + 轻量化设计]").font.size = Pt(9)
    doc.add_page_break()

    # ---- 0. 摘要 + 硬指标达标对照表 ----
    doc.add_heading("0. 摘要 · 硬指标达标对照表", level=1)
    doc.add_paragraph(
        "本方案面向无人机/飞救生艇低空俯拍场景下的海上落水搜救,目标:在机载边缘端实现"
        "高精度(各类 PR 曲线)且实时(动态视频 ≥30FPS)的端到端视觉识别,识别落水人员/船只/浮标 3 类。"
        "写作原则:开篇即以硬指标对照表焊死出局线,后续每个技术选择回答'为什么—对上赛题哪条约束'。")
    _add_table(doc,
               ["硬指标", "官方要求", "本方案达成", "证据"],
               [["边缘端实时性", "动态视频 ≥30FPS", "___ FPS(含编码端到端)", "性能报告 §3 三档 FPS"],
                ["检测精度", "各类 PR 曲线", "落水人员/船只/浮标", "性能报告 §2"],
                ["平台", "Jetson 等机载边缘", "Jetson Orin Nano 8G", "本方案 §5 部署"]],
               col_widths=[1.3, 1.6, 1.8, 1.7])

    # ---- 1. 问题分析与赛题理解 ----
    doc.add_heading("1. 问题分析与赛题理解", level=1)
    doc.add_paragraph("业务场景:无人机 / 飞救生艇视频感知,海上落水人员搜救,要求发现目标并给出可航行坐标。")
    doc.add_paragraph("三大痛点(对应官方要点):")
    for s in ["① 大范围监测中,中小目标(落水人员常仅 6–16px)识别难;",
              "② 水面光照反射 / 波纹 / 碎浪亮点造成强干扰,易产生误检(false positive);",
              "③ 机载算力受限,模型须在 Orin Nano 8G 上实时(≥30FPS)运行。"]:
        doc.add_paragraph(s, style="List Bullet")
    doc.add_paragraph("目标:高精度 + 实时(≥30FPS)端到端视觉识别系统,识别落水人员 / 船只 / 浮标 3 类。")

    # ---- 2. 模型选型依据 ----
    doc.add_heading("2. 模型选型依据(官方要点①)", level=1)
    doc.add_paragraph(
        "选用 YOLOv12(官方推荐前沿架构)+ P2 小目标检测头,规模档位 scale=n(端侧友好)。"
        "选型逻辑均围绕 30FPS 出局线与小目标召回展开:")
    _add_table(doc,
               ["选择", "为什么(对上的约束)", "代价 / 取舍"],
               [["YOLOv12-n 而非 -s", "Orin Nano 上 n→s 掉的 FPS 比 640→768 更贵,先保过线", "精度上界用 s 仅作对照,不上端侧"],
                ["+ P2(stride-4)头", "对 6–16px 小目标最大单一杠杆,直接抬 tiny/small 召回", "增高分辨率层,端侧延迟上升,需 INT8/分辨率管控补偿"],
                ["不上实时 SAHI 切片", "实时切片与 30FPS 死冲", "SAHI 仅作离线高召回对照档"],
                ["内置 area-attention", "选 YOLOv12 即满足'注意力增强'要求", "不再额外叠 CBAM(避免延迟与冗余)"]],
               col_widths=[1.6, 2.6, 2.1])
    doc.add_paragraph("必做消融:四头 vs 去 P5 三头(海面几乎无大目标,P5 可能冗余,砍掉换帧率)——见 §8 / 性能报告 §4。")

    # ---- 3. 海上小目标特征提取 ----
    doc.add_heading("3. 海上小目标特征提取(官方要点②)", level=1)
    doc.add_paragraph("以高分辨率检测头 + 小目标友好的标签分配/损失为主线:")
    for s in ["P2 高分辨率检测头(stride-4)+ 训练分辨率 1024,保住小目标空间信息;",
              "标签分配/损失:NWD(归一化 Wasserstein 距离,对小目标中心位置偏移不敏感)+ Wise-IoU / Inner-IoU;",
              "每项增益以本队自测消融为准(引用思路来源,不照搬论文百分比)。"]:
        doc.add_paragraph(s, style="List Bullet")
    _maybe_add_figure(doc, "report_fig1_bucket_recall.png",
                      "图3-1 按目标像素面积分桶的召回(占位值)。P2+NWD+GT-Glint 链路在 tiny/small 带来最大增益,证明小目标不漏检。")

    # ---- 4. 注意力机制增强 + 反光误报抑制 ----
    doc.add_heading("4. 注意力机制增强与水面反光误报抑制(官方要点③)", level=1)
    doc.add_paragraph(
        "注意力:采用 YOLOv12 内置 area-attention 满足官方'注意力增强'要求;在 30FPS 约束下不额外叠 CBAM。")
    doc.add_paragraph(
        "水面反光误报抑制链(本方案核心算法创新):P2 + area-attention + NWD 标签分配 + "
        "GT-Anchored Glint 难负样本。其中 GT-Anchored Glint 在远离 GT 框的水面区域贴高斯高光斑、"
        "标签保持不变,作为背景难负样本喂入,逼模型学习'亮 ≠ 目标',显著抑制反光误检。")
    _maybe_add_figure(doc, "augment_water_before_after_0.png",
                      "图4-1 GT-Anchored Glint 增广 before/after(合成海面演示)。高光只贴在远离 GT 的水面区域;标签不变。")

    # ---- 5. 轻量化网络与端侧部署 ----
    doc.add_heading("5. 轻量化网络设计与端侧部署(官方要点④)", level=1)
    for s in ["规模:yolo12n + TensorRT INT8 混合精度(对小目标敏感的层回退 FP16);",
              "校准:用真实海况帧(反光/波纹/运动模糊 300–500 张)做 INT8 校准,禁用干净 val 图;",
              "功耗:Super Mode(JetPack 6.2,nvpmodel -m 2,MAXN_SUPER 25W)+ jetson_clocks,零成本约 1.7x 吞吐,DC 供电;",
              "量化纪律:每改一处(P2 / 分辨率 / INT8)都重测分桶召回,掉 >2% AP_small 即退 FP16,禁止只看总 mAP。"]:
        doc.add_paragraph(s, style="List Bullet")
    doc.add_paragraph("降级链(精度优先级低→高,按序砍以守住 30FPS):"
                      "768→640 → 关 P2/砍 P5 → INT8 替 FP16 → 确认 Super Mode + DC 供电 → 抽帧推理 → 最后才动模型。")

    # ---- 6. 系统设计与闭环 ----
    doc.add_heading("6. 系统设计与闭环(创新 / 加分)", level=1)
    doc.add_paragraph(
        "端到端闭环:端侧检测 → 检测框→GPS 救援航点(针孔模型 + 海平面 z=0 求交)→ "
        "地面站可视化决策(QGroundControl 视频流融合)。架构论证:检测在边缘、地面站只做可视化,"
        "是边缘计算 + 低延迟的正确架构。时序一致性滤波(连续 ≥k 帧才确认)进一步抑制反光闪点误报。")
    _maybe_add_figure(doc, "report_fig3_fps_accuracy.png",
                      "图6-1 FPS-精度权衡 + 30FPS 红线 + 帕累托前沿(占位值),用于选定端侧最优工作点。")

    # ---- 7. 三个创新点 ----
    doc.add_heading("7. 三个创新点(对应答辩)", level=1)
    _add_table(doc,
               ["#", "创新点", "对评分维度的贡献"],
               [["1", "检测框→GPS 救援航点系统闭环(诚实报投影误差区间,几十米级)",
                 "系统创新,直击命题方'发现落水者并给可航行坐标'业务语言"],
                ["2", "水面反光误报抑制链(P2+area-attn+NWD+GT-Anchored Glint),精度-速度双档",
                 "算法创新,消融表论证每项独立增益"],
                ["3", "Orin Nano 无硬件编码器约束下的端到端实时管线工程优化",
                 "部署创新,工程严谨性 = 综合表现力得分点"]],
               col_widths=[0.4, 3.6, 2.7])

    # ---- 8. 实验与结论(指向性能报告) ----
    doc.add_heading("8. 实验与结论(指向性能报告)", level=1)
    doc.add_paragraph("数据:SeaDronesSee + AFO 公开底座 + 命题方二开数据(域适配微调)。完整实验见性能评估报告,核心消融表如下(真值待回填):")
    _add_table(doc,
               ["配置", "mAP", "small 召回", "端侧 FPS", "结论"],
               [["baseline (yolo12n)", "___", "___", "___", "基线"],
                ["+ P2 头", "___", "___", "___", "小目标↑"],
                ["+ NWD 标签分配", "___", "___", "___", "位置鲁棒↑"],
                ["+ Wise-IoU", "___", "___", "___", "回归↑"],
                ["+ 物理增广 + GT-Glint 难负样本", "___", "___", "___", "反光误报↓"],
                ["去 P5(四头→三头)", "___", "___", "___", "帧率↑ / 精度?"],
                ["+ INT8 量化", "___", "___", "___", "速度↑ / 掉点?"]],
               col_widths=[2.6, 0.9, 1.1, 1.0, 1.2])
    doc.add_paragraph("消融纪律:每行只改一项,其余固定,禁止一次叠多项凭感觉。")
    _maybe_add_figure(doc, "report_fig2_pr_curves.png",
                      "图8-1 各类 Precision-Recall 曲线(占位值)。救援场景 recall 优先,落水人员高 recall 区为验收重点。")
    doc.add_heading("局限与展望", level=2)
    doc.add_paragraph(
        "局限:命题方二开数据存在域差,类别口径(穿救生衣落水者是否独立判定)需拿到数据后核对;"
        "INT8 对小目标敏感;Orin Nano 无硬件 NVENC,软编码 H265 占 CPU。"
        "展望:伪标签 UDA 域适配微调、DeepStream 流水线加速、更稳健的时序跟踪。")

    return doc


def generate(out_dir=OUT_DIR):
    out_dir = Path(out_dir); out_dir.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    out_path = out_dir / "技术方案_初稿.docx"
    doc.save(str(out_path))
    return out_path


def _selftest():
    import tempfile
    from docx import Document
    ok = True

    def check(c, m):
        nonlocal ok
        print(("  OK  " if c else "  XX  ") + m)
        ok = ok and c

    with tempfile.TemporaryDirectory() as td:
        out = generate(Path(td))
        check(out.exists() and out.stat().st_size > 8000, f"docx 生成且非空({out.stat().st_size}B)")
        # 重新打开校验结构
        d = Document(str(out))
        headings = [p.text for p in d.paragraphs if p.style.name.startswith("Heading")]
        check(len(headings) >= 9, f"章节标题数 {len(headings)} (含 0-8 + 子节)")
        for key in ["选型", "小目标", "注意力", "轻量化", "创新"]:
            check(any(key in h for h in headings), f"含章节关键词「{key}」")
        check(len(d.tables) >= 4, f"含表格 {len(d.tables)} 张(达标对照/选型/创新/消融)")
        # 校验硬指标对照表表头
        t0 = d.tables[0]
        hdr = [c.text for c in t0.rows[0].cells]
        check("硬指标" in hdr and "证据" in hdr, f"首表为硬指标达标对照表({hdr})")
        # 校验消融表含 7 行配置
        ablation = [t for t in d.tables if t.rows[0].cells[0].text == "配置"]
        check(len(ablation) >= 1 and len(ablation[0].rows) >= 8, "消融表含 >=7 行配置")

    print("\n" + ("OK 技术方案 docx 生成器 自测通过" if ok else "XX 自测未通过"))
    return ok


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--selftest", action="store_true")
    ap.add_argument("--out", default=str(OUT_DIR))
    a = ap.parse_args()
    if a.selftest:
        sys.exit(0 if _selftest() else 1)
    out = generate(a.out)
    figs_in = sum(1 for f in ["report_fig1_bucket_recall.png", "augment_water_before_after_0.png",
                              "report_fig3_fps_accuracy.png", "report_fig2_pr_curves.png"]
                  if (FIG_DIR / f).exists())
    print(f"[OK] 生成 {out}  ({os.path.getsize(out)} bytes, 内嵌 {figs_in}/4 张配图)")


if __name__ == "__main__":
    main()
