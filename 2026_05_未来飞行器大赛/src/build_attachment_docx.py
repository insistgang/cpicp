# -*- coding: utf-8 -*-
"""
附件1 技术成熟度报告(TRL) Word 渲染器 · 薄封装

复用 build_report_docx.py 的通用 Markdown 解析 + 渲染管线:
  parse_markdown / clean_inline / render_table / add_heading /
  add_body_paragraph / add_mono_block / setup_page / 字体工具

与正文报告书的差异(故单独成脚本而非共用 build()):
  - 附件无配图锚点(纯文字+表格+TRL演进 ASCII 框图)
  - 封面是"附件1"样式, 不重复正文报告书的赛事封面/目录域
  - 自测校验的关键文本是 TRL 报告自己的章节, 与正文不同

输入: ../技术成熟度报告_附件1.md
输出: ../output/技术成熟度报告_附件1.docx

运行:  python3 build_attachment_docx.py
       python3 build_attachment_docx.py --input <md> --output <docx>
自测:  python3 build_attachment_docx.py --selftest
"""
import os
import sys
import argparse

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 复用正文渲染器的通用管线(同目录)
HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)
import build_report_docx as B

ROOT = os.path.abspath(os.path.join(HERE, ".."))
MD_PATH = os.path.join(ROOT, "技术成熟度报告_附件1.md")
OUT_DIR = os.path.join(ROOT, "output")
OUT_DOCX = os.path.join(OUT_DIR, "技术成熟度报告_附件1.docx")

# 自测期望出现的关键文本(取自附件1自身章节)
EXPECT_KEYWORDS = [
    "技术成熟度报告",
    "TRL 等级总览",
    "技术底座详细说明",
    "技术风险与缓解措施",
    "技术演进路线",
    "与现有技术对比",
]


def add_attachment_cover(doc):
    """附件封面: 简洁标识 + 副标题, 不重复正文报告书的赛事封面/目录域。"""
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("附件 1")
    B.set_run_font(r, B.FONT_HEAD, 22, bold=True, color=(0x1F, 0x3A, 0x6B))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("技术成熟度报告（TRL）")
    B.set_run_font(r, B.FONT_HEAD, 18, bold=True, color=(0x1F, 0x3A, 0x6B))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基于低空智能感知的城市违建巡查系统")
    B.set_run_font(r, B.FONT_BODY, 13, color=(0x44, 0x44, 0x44))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("——技术底座成熟度评估 · 证明方案可实现性")
    B.set_run_font(r, B.FONT_BODY, 11, color=(0x77, 0x77, 0x77))

    doc.add_page_break()


def build(md_path=MD_PATH, out_path=OUT_DOCX):
    with open(md_path, "r", encoding="utf-8") as f:
        md = f.read()
    blocks = B.parse_markdown(md)

    doc = Document()
    # 复用正文渲染器一致的 Normal 中文字体设置
    normal = doc.styles["Normal"]
    normal.font.name = B.FONT_BODY
    normal.font.size = Pt(12)
    normal._element.get_or_add_rPr().append(B._mk_rfonts(B.FONT_BODY))

    B.setup_page(doc)
    add_attachment_cover(doc)

    stats = {"headings": 0, "paras": 0, "tables": 0,
             "lists": 0, "code": 0}
    first_h1_skipped = False

    for blk in blocks:
        t = blk["type"]
        if t == "heading":
            level = blk["level"]
            text = B.clean_inline(blk["text"])
            # 首个一级标题(# 技术成熟度报告...)已在封面表达, 跳过避免重复
            if level == 1 and not first_h1_skipped:
                first_h1_skipped = True
                continue
            B.add_heading(doc, text, level)
            stats["headings"] += 1
        elif t == "para":
            B.add_body_paragraph(doc, B.clean_inline(blk["text"]))
            stats["paras"] += 1
        elif t == "blockquote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(B.clean_inline(" ".join(blk["lines"])))
            B.set_run_font(run, B.FONT_BODY, 10, color=(0x77, 0x77, 0x77))
            stats["paras"] += 1
        elif t == "table":
            B.render_table(doc, blk["rows"])
            stats["tables"] += 1
        elif t == "code":
            B.add_mono_block(doc, blk["lines"])
            stats["code"] += 1
        elif t == "list":
            for it in blk["items"]:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(B.clean_inline(it["text"]))
                B.set_run_font(run, B.FONT_BODY, 11)
            stats["lists"] += 1
        elif t == "hr":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    return out_path, stats


def selftest(md_path=MD_PATH, out_path=OUT_DOCX):
    out_path, stats = build(md_path, out_path)
    print(f"[OK] 生成 {os.path.basename(out_path)}  "
          f"({os.path.getsize(out_path)} bytes)")
    ok = True
    if not os.path.exists(out_path) or os.path.getsize(out_path) < 5000:
        print("[FAIL] docx 不存在或过小")
        return False

    # 读回校验: 确认 python-docx 能重新打开(合法 OOXML)并统计
    d = Document(out_path)
    n_para = len(d.paragraphs)
    n_tbl = len(d.tables)
    n_nonempty = sum(1 for p in d.paragraphs if p.text.strip())
    n_head = sum(1 for p in d.paragraphs
                 if p.style is not None
                 and str(p.style.name).startswith("Heading"))

    print(f"[读回] 段落总数={n_para}  非空段落={n_nonempty}  "
          f"标题段落={n_head}  表格={n_tbl}")
    print(f"[渲染统计] {stats}")

    checks = [
        (n_para >= 30, f"段落数应≥30, 实际{n_para}"),
        (n_nonempty >= 25, f"非空段落应≥25, 实际{n_nonempty}"),
        (n_head >= 8, f"标题段落应≥8, 实际{n_head}"),
        (stats["headings"] >= 8, f"渲染标题数应≥8, 实际{stats['headings']}"),
        (n_tbl >= 4, f"表格数应≥4, 实际{n_tbl}"),
        (stats["code"] >= 1, f"代码/框图块应≥1, 实际{stats['code']}"),
    ]
    for cond, msg in checks:
        print(("[PASS] " if cond else "[FAIL] ") + msg)
        ok = ok and cond

    # 校验关键章节文本存在
    all_text = "\n".join(p.text for p in d.paragraphs)
    # 表格内文本也纳入(TRL 等级在表格里)
    for tb in d.tables:
        for row in tb.rows:
            for cell in row.cells:
                all_text += "\n" + cell.text
    for kw in EXPECT_KEYWORDS:
        present = kw in all_text
        print(("[PASS] " if present else "[FAIL] ") + f"含关键文本: {kw}")
        ok = ok and present
    # TRL 表格内容应被保留
    trl_present = ("TRL 9" in all_text) and ("综合 TRL" in all_text)
    print(("[PASS] " if trl_present else "[FAIL] ") + "保留 TRL 表格内容(TRL 9 / 综合 TRL)")
    ok = ok and trl_present

    print("=== 附件1 DOCX SELFTEST", "PASS ===" if ok else "FAIL ===")
    return ok


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--selftest", action="store_true")
    ap.add_argument("--input", default=MD_PATH, help="输入 Markdown 路径")
    ap.add_argument("--output", default=OUT_DOCX, help="输出 docx 路径")
    args = ap.parse_args()
    if args.selftest:
        sys.exit(0 if selftest(args.input, args.output) else 1)
    else:
        out, stats = build(args.input, args.output)
        print(f"已生成: {out}")
        print(f"统计: {stats}")
