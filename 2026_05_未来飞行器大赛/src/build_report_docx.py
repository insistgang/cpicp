# -*- coding: utf-8 -*-
"""
项目报告书 Word 渲染器 · 低空智能感知城市违建巡查系统

输入: ../项目报告书_初稿.md  (7 章 Markdown)
依据: docs/...项目报告书(赛道1.3.4)参考模板.txt 的官方 5 段骨架
输出: ../output/项目报告书_低空违建巡查系统.docx

实现:
  - 封面页(参赛队名称/单位/队员/联系人/日期, 对齐官方模板字段)
  - 目录占位说明(Word 打开后可一键更新域 / 提交前转 PDF)
  - 章节标题层级(# -> Heading1, ## -> Heading2, ### -> Heading3 ...)
  - Markdown 表格 -> Word 原生表格(带表头底纹)
  - 代码块/ASCII 框图 -> 等宽字体段落(并在对应位置插入 matplotlib 生成的 PNG 配图)
  - 4 张 PNG 图按章节锚点插入(架构图/流程图/商业图/交叉矩阵)
  - 页边距 A4 标准(上下2.54cm 左右3.17cm)、正文宋体小四、标题黑体

纯 python-docx(已装), 离线可跑。
运行:  python3 build_report_docx.py
自测:  python3 build_report_docx.py --selftest   (生成后用 python-docx 读回校验段落/表格/图片数)
"""
import os
import re
import sys
import argparse

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, ".."))
MD_PATH = os.path.join(ROOT, "项目报告书_初稿.md")
FIG_DIR = os.path.join(ROOT, "output", "figures")
OUT_DIR = os.path.join(ROOT, "output")
OUT_DOCX = os.path.join(OUT_DIR, "项目报告书_低空违建巡查系统.docx")

# 正文/标题字体
FONT_BODY = "宋体"
FONT_HEAD = "黑体"
FONT_MONO = "Consolas"

# 图片锚点: 当解析到某个二级/三级标题文本包含关键字时, 在该节末尾插入对应图
# (key 用于匹配标题, value 为 (png文件名, 图说明))
FIG_ANCHORS = [
    ("2.1", "系统总体架构", "fig_architecture.png", "图1 端-边-云三层系统总体架构"),
    ("2.4", "实施流程", "fig_workflow.png", "图2 系统实施流程(四阶段)"),
    ("4.2", "商业模式", "fig_business.png", "图3 商业模式与成本收益"),
    ("5", "交叉融合", "fig_crossmatrix.png", "图4 六学科交叉融合矩阵"),
]


# ---------------------------------------------------------------------------
# 样式工具
# ---------------------------------------------------------------------------
def set_run_font(run, name=FONT_BODY, size=12, bold=False, color=None,
                 east_asian=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    # 设置中文字体(east asian), 否则中文会回退到默认字体
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asian or name)


def set_cell_bg(cell, hex_color):
    """给单元格设置底纹颜色(表头用)。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_body_paragraph(doc, text, size=12, indent_chars=2):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Pt(size * indent_chars)  # 首行缩进2字符
    pf.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, FONT_BODY, size)
    return p


def add_heading(doc, text, level):
    sizes = {1: 16, 2: 14, 3: 13, 4: 12}
    size = sizes.get(level, 12)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.style = doc.styles[f"Heading {min(level,4)}"]
    run = p.add_run(text)
    set_run_font(run, FONT_HEAD, size, bold=True, color=(0x1F, 0x3A, 0x6B))
    return p


def add_mono_block(doc, lines):
    """ASCII 框图 / 代码块: 等宽字体, 无缩进, 灰底。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run("\n".join(lines))
    set_run_font(run, FONT_MONO, 9, east_asian=FONT_BODY)
    return p


def add_figure(doc, png_path, caption, width_inch=6.0):
    if not os.path.exists(png_path):
        print(f"[WARN] 缺图, 跳过插入: {png_path}", file=sys.stderr)
        return False
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(png_path, width=Inches(width_inch))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    crun = cap.add_run(caption)
    set_run_font(crun, FONT_BODY, 10.5, bold=True, color=(0x44, 0x44, 0x44))
    cap.paragraph_format.space_after = Pt(8)
    return True


# ---------------------------------------------------------------------------
# Markdown 轻量解析(只覆盖本报告用到的语法: 标题/段落/表格/代码块/列表)
# ---------------------------------------------------------------------------
def parse_markdown(md_text):
    """返回 block 列表, 每个 block 是 dict: {type, ...}。
    type ∈ heading|para|table|code|list|hr|blockquote
    """
    lines = md_text.splitlines()
    blocks = []
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 代码块 ```
        if stripped.startswith("```"):
            j = i + 1
            code = []
            while j < n and not lines[j].strip().startswith("```"):
                code.append(lines[j])
                j += 1
            blocks.append({"type": "code", "lines": code})
            i = j + 1
            continue

        # 水平线
        if re.match(r"^-{3,}$", stripped):
            blocks.append({"type": "hr"})
            i += 1
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            blocks.append({"type": "heading", "level": level,
                           "text": m.group(2).strip()})
            i += 1
            continue

        # 引用块 (> ...)
        if stripped.startswith(">"):
            quote = []
            while i < n and lines[i].strip().startswith(">"):
                quote.append(lines[i].strip().lstrip(">").strip())
                i += 1
            blocks.append({"type": "blockquote", "lines": quote})
            continue

        # 表格 (连续含 | 的行, 且第二行是分隔行)
        if "|" in line and i + 1 < n and re.match(r"^\s*\|?[\s:|-]+\|?\s*$",
                                                  lines[i + 1]):
            tbl = []
            while i < n and "|" in lines[i] and lines[i].strip():
                tbl.append(lines[i])
                i += 1
            rows = []
            for ridx, r in enumerate(tbl):
                if ridx == 1:  # 分隔行跳过
                    continue
                cells = [c.strip() for c in r.strip().strip("|").split("|")]
                rows.append(cells)
            blocks.append({"type": "table", "rows": rows})
            continue

        # 列表(- 或 数字.)
        if re.match(r"^\s*([-*]|\d+\.)\s+", line):
            items = []
            while i < n and re.match(r"^\s*([-*]|\d+\.)\s+", lines[i]):
                txt = re.sub(r"^\s*([-*]|\d+\.)\s+", "", lines[i])
                indent = len(lines[i]) - len(lines[i].lstrip())
                items.append({"text": txt.strip(), "indent": indent})
                i += 1
            blocks.append({"type": "list", "items": items})
            continue

        # 空行
        if not stripped:
            i += 1
            continue

        # 普通段落(合并连续非空、非特殊行)
        para = [stripped]
        i += 1
        while i < n:
            nxt = lines[i].strip()
            if (not nxt or nxt.startswith("#") or nxt.startswith("```")
                    or nxt.startswith(">") or re.match(r"^-{3,}$", nxt)
                    or ("|" in lines[i])
                    or re.match(r"^\s*([-*]|\d+\.)\s+", lines[i])):
                break
            para.append(nxt)
            i += 1
        blocks.append({"type": "para", "text": " ".join(para)})
    return blocks


def clean_inline(text):
    """去掉行内 markdown 强调标记 ** * `, 保留纯文本(Word 里不再用 md 语法)。"""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[([ xX]?)\]", lambda m: "☑" if m.group(1).strip() else "☐",
                  text)
    return text


# ---------------------------------------------------------------------------
# 封面 + 页面设置
# ---------------------------------------------------------------------------
def setup_page(doc):
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)


def add_cover(doc):
    # 顶部赛事名
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("第十二届中国研究生未来飞行器创新大赛")
    set_run_font(r, FONT_HEAD, 18, bold=True, color=(0x1F, 0x3A, 0x6B))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("参赛作品项目报告书")
    set_run_font(r, FONT_HEAD, 16, bold=True, color=(0x1F, 0x3A, 0x6B))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("(交叉赛道 3.1 · 低空未来飞行器应用场景设计)")
    set_run_font(r, FONT_BODY, 12, color=(0x66, 0x66, 0x66))

    for _ in range(2):
        doc.add_paragraph()

    # 作品名称
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基于低空智能感知的城市违建巡查系统")
    set_run_font(r, FONT_HEAD, 17, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("——低空经济在公共治理的落地探索")
    set_run_font(r, FONT_BODY, 13, color=(0x44, 0x44, 0x44))

    for _ in range(3):
        doc.add_paragraph()

    # 信息字段表(对齐官方模板字段, 蓝色占位待填)
    fields = [
        ("参赛队名称", "（待填写）"),
        ("参赛队单位", "（待填写，提交PDF前移除并匿名）"),
        ("参赛队其他单位", "（如有）"),
        ("参赛队员", "（待填写，建议 8 人，含城乡规划/公共管理/GIS 方向）"),
        ("联系人 / 邮箱 / 电话", "（待填写）"),
    ]
    tbl = doc.add_table(rows=len(fields), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for ri, (k, v) in enumerate(fields):
        c0 = tbl.cell(ri, 0)
        c1 = tbl.cell(ri, 1)
        c0.width = Cm(4.5)
        c1.width = Cm(9.5)
        set_cell_bg(c0, "EAF0FB")
        pr = c0.paragraphs[0]
        run = pr.add_run(k)
        set_run_font(run, FONT_HEAD, 11, bold=True)
        pr2 = c1.paragraphs[0]
        run2 = pr2.add_run(v)
        # 蓝色提示文字, 对齐官方"移除所有蓝色文字"约定
        set_run_font(run2, FONT_BODY, 11, color=(0x2E, 0x5A, 0xAC))

    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("日期：      年      月      日")
    set_run_font(r, FONT_BODY, 12)

    # 重要提示(蓝色, 提交前删除)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("重要提示：1.提交前请移除全部蓝色提示文字；2.以PDF格式提交；"
                  "3.全文匿名，不得出现校名/导师/校徽")
    set_run_font(r, FONT_BODY, 9, color=(0x2E, 0x5A, 0xAC))

    doc.add_page_break()


def add_toc_placeholder(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("目  录")
    set_run_font(r, FONT_HEAD, 15, bold=True)
    # 插入 Word 域代码 TOC, 打开文档后右键"更新域"即可生成真实目录
    para = doc.add_paragraph()
    run = para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_txt = OxmlElement("w:t")
    fld_txt.text = "（在 Word 中右键此处选择“更新域”以生成目录）"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr, fld_sep, fld_txt, fld_end):
        run._r.append(el)
    doc.add_page_break()


# ---------------------------------------------------------------------------
# 主渲染
# ---------------------------------------------------------------------------
def render_table(doc, rows):
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncol)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        for ci in range(ncol):
            cell = tbl.cell(ri, ci)
            txt = clean_inline(row[ci]) if ci < len(row) else ""
            para = cell.paragraphs[0]
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            run = para.add_run(txt)
            if ri == 0:
                set_run_font(run, FONT_HEAD, 10.5, bold=True)
                set_cell_bg(cell, "1F3A6B")
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            else:
                set_run_font(run, FONT_BODY, 10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def maybe_insert_figure(doc, heading_text):
    """根据当前标题文本判断是否插入对应配图。返回插入的图说明列表。"""
    inserted = []
    for sec_no, kw, png, cap in FIG_ANCHORS:
        if sec_no in heading_text and kw in heading_text:
            ok = add_figure(doc, os.path.join(FIG_DIR, png), cap)
            if ok:
                inserted.append(cap)
    return inserted


def build(md_path=MD_PATH, out_path=OUT_DOCX):
    with open(md_path, "r", encoding="utf-8") as f:
        md = f.read()
    blocks = parse_markdown(md)

    doc = Document()
    # 默认正文样式中文字体
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(12)
    normal._element.get_or_add_rPr().append(
        _mk_rfonts(FONT_BODY))

    setup_page(doc)
    add_cover(doc)
    add_toc_placeholder(doc)

    stats = {"headings": 0, "paras": 0, "tables": 0, "figures": 0,
             "lists": 0, "code": 0}
    pending_heading_text = ""

    # 跳过初稿正文最顶部的标题(已在封面体现): 找到第一个 "## 1." 之前的引用块/一级标题不重复渲染封面
    first_h1_skipped = False

    for blk in blocks:
        t = blk["type"]
        if t == "heading":
            level = blk["level"]
            text = clean_inline(blk["text"])
            # 初稿的一级标题(# 项目报告书初稿...)在封面已表达, 跳过首个
            if level == 1 and not first_h1_skipped:
                first_h1_skipped = True
                continue
            add_heading(doc, text, level)
            stats["headings"] += 1
            pending_heading_text = text
            ins = maybe_insert_figure(doc, text)
            stats["figures"] += len(ins)
        elif t == "para":
            add_body_paragraph(doc, clean_inline(blk["text"]))
            stats["paras"] += 1
        elif t == "blockquote":
            # 引用块用浅灰小字呈现(初稿里多为元信息/说明)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(clean_inline(" ".join(blk["lines"])))
            set_run_font(run, FONT_BODY, 10, color=(0x77, 0x77, 0x77))
            stats["paras"] += 1
        elif t == "table":
            render_table(doc, blk["rows"])
            stats["tables"] += 1
        elif t == "code":
            add_mono_block(doc, blk["lines"])
            stats["code"] += 1
        elif t == "list":
            for it in blk["items"]:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.line_spacing_rule = \
                    WD_LINE_SPACING.ONE_POINT_FIVE
                run = p.add_run(clean_inline(it["text"]))
                set_run_font(run, FONT_BODY, 11)
            stats["lists"] += 1
        elif t == "hr":
            # 分隔线渲染为一条浅色短段落分隔
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

    # 文末: 若某些锚点图因标题未精确命中而漏插, 兜底集中插入在"附件"前
    inserted_caps = _collect_inserted(doc)
    missing = [(png, cap) for _, _, png, cap in FIG_ANCHORS
               if cap not in inserted_caps]
    if missing:
        add_heading(doc, "附图汇总", 2)
        for png, cap in missing:
            if add_figure(doc, os.path.join(FIG_DIR, png), cap):
                stats["figures"] += 1

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    return out_path, stats


def _mk_rfonts(name):
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    return rfonts


def _collect_inserted(doc):
    """统计已插入图片的图说明(用文末兜底判断)。简单实现: 扫描文档段落文本。"""
    caps = set()
    for p in doc.paragraphs:
        txt = p.text.strip()
        for _, _, _, cap in FIG_ANCHORS:
            if txt == cap:
                caps.add(cap)
    return caps


# ---------------------------------------------------------------------------
# 自测: 读回校验
# ---------------------------------------------------------------------------
def selftest():
    out_path, stats = build()
    print(f"[OK] 生成 {os.path.basename(out_path)}  "
          f"({os.path.getsize(out_path)} bytes)")
    ok = True
    if not os.path.exists(out_path) or os.path.getsize(out_path) < 10000:
        print("[FAIL] docx 不存在或过小"); return False

    # 读回校验
    d = Document(out_path)
    n_para = len(d.paragraphs)
    n_tbl = len(d.tables)
    # 统计内嵌图片: 文档关系里 image 类型数量
    n_img = 0
    for rel in d.part.rels.values():
        if "image" in rel.reltype:
            n_img += 1
    # 非空段落数
    n_nonempty = sum(1 for p in d.paragraphs if p.text.strip())

    print(f"[读回] 段落总数={n_para}  非空段落={n_nonempty}  "
          f"表格={n_tbl}  内嵌图片={n_img}")
    print(f"[渲染统计] {stats}")

    checks = [
        (n_para >= 50, f"段落数应≥50, 实际{n_para}"),
        (n_nonempty >= 40, f"非空段落应≥40, 实际{n_nonempty}"),
        (n_tbl >= 8, f"表格数应≥8, 实际{n_tbl}"),
        (n_img >= 4, f"内嵌图片应≥4, 实际{n_img}"),
        (stats["figures"] >= 4, f"插图统计应≥4, 实际{stats['figures']}"),
    ]
    for cond, msg in checks:
        print(("[PASS] " if cond else "[FAIL] ") + msg)
        ok = ok and cond

    # 校验封面关键字 + 章节标题存在
    all_text = "\n".join(p.text for p in d.paragraphs)
    for kw in ["第十二届中国研究生未来飞行器创新大赛", "基于低空智能感知的城市违建巡查系统",
               "设计理念", "作品设计", "创新点", "实施与保障", "交叉融合"]:
        present = kw in all_text
        print(("[PASS] " if present else "[FAIL] ") + f"含关键文本: {kw}")
        ok = ok and present

    print("=== DOCX SELFTEST", "PASS ===" if ok else "FAIL ===")
    return ok


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--selftest", action="store_true")
    args = ap.parse_args()
    if args.selftest:
        sys.exit(0 if selftest() else 1)
    else:
        out, stats = build()
        print(f"已生成: {out}")
        print(f"统计: {stats}")
